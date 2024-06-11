from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('Mahin Malhotra', 0)

# Contact Information
doc.add_paragraph(
    'Python Developer\n'
    'LinkedIn: https://in.linkedin.com/in/mahinm20\n'
    'Email: mahinmalhotra20@gmail.com\n'
    'Phone: +91-7696464478'
)

# Add a horizontal line
doc.add_paragraph('')

# Summary
doc.add_heading('Summary', level=1)
doc.add_paragraph(
    "Experienced Software Developer in Test (SDET) with over 5 years of experience in developing automation frameworks and optimizing CI/CD pipelines. "
    "Adept at using Python, Django, Pytest, and Playwright to improve software quality and efficiency. Proven ability to lead projects, collaborate with cross-functional teams, "
    "and enhance code quality. Recognized for outstanding performance with multiple awards."
)

# Skills
doc.add_heading('Skills', level=1)
doc.add_paragraph(
    'Programming Languages: Python\n'
    'Frameworks: Django, Pytest, Playwright\n'
    'CI/CD Tools: Jenkins, Gitlab CI/CD\n'
    'Other Tools: Git, Docker, Kafka, Lens, SonarScan\n'
    'Database: PostgresSQL'
)

# Professional Experience
doc.add_heading('Professional Experience', level=1)

# Job 1
doc.add_heading('Software Developer in Test', level=2)
doc.add_paragraph('Magic Finserv | Noida, India | Sep 2022 – Present')
doc.add_paragraph(
    '- Led the development of a Python-based automation test framework for a UK-based client, reducing UI and backend bugs pre-production.\n'
    '- Managed and executed over 1000 test cases using Zephyr, streamlining organization and tracking.\n'
    '- Optimized CI/CD pipelines on Gitlab, reducing regression testing time by 75%.\n'
    '- Collaborated with onshore QA teams for sprint planning and test strategy, ensuring comprehensive test coverage.\n'
    '- Served as a key code reviewer, enhancing code quality across multiple projects.\n'
    '- Instrumental in the development and release of a GPT-4 powered web app for an esteemed client, utilizing AI for test case generation and reporting.'
)

# Job 2
doc.add_heading('Release and Automation Engineer', level=2)
doc.add_paragraph('Wipro Technologies | Bengaluru, India | Aug 2018 – Jul 2022')
doc.add_paragraph(
    '- Implemented an end-to-end Jenkins project, enabling blue/green deployments for over 200 applications.\n'
    '- Developed Python automation scripts that reduced manual effort from 1 hour to 20 minutes.\n'
    '- Handled daily build release activities including creating Jenkins CI/CD pipelines, Git repositories, and UCD components.\n'
    '- Wrote a Python script to delete over 1000 stale git flow branches, maintaining repository cleanliness.'
)

# Education
doc.add_heading('Education', level=1)
doc.add_paragraph(
    'B. Tech, Computer Science\n'
    'Chitkara University | Punjab, India | Aug 2014 – Jul 2018\n'
    'CGPA: 7.33'
)

# Awards
doc.add_heading('Awards', level=1)
doc.add_paragraph(
    '- Role Model Award, Magic Finserv | Jun 2023\n'
    '  Awarded for exceptional performance and leadership in automation projects.\n'
    '- Rising Star Award, Magic Finserv | Jan 2023\n'
    '  Recognized as the best performing newcomer for significant contributions to the team.'
)

# Save the document
doc_path = "Mahin_Malhotra_Improved_Resume.docx"
doc.save(doc_path)
doc_path
